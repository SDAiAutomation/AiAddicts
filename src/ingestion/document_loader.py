import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}


def load_pdf(file_path: str) -> tuple[str, dict]:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages_text = []
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            pages_text.append(extracted)

    return "\n".join(pages_text), {
        "source": Path(file_path).name,
        "file_path": str(file_path),
        "file_type": "pdf",
        "page_count": len(reader.pages),
    }


def load_docx(file_path: str) -> tuple[str, dict]:
    from docx import Document

    doc = Document(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts), {
        "source": Path(file_path).name,
        "file_path": str(file_path),
        "file_type": "docx",
        "page_count": None,
    }


def load_txt(file_path: str) -> tuple[str, dict]:
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    return text, {
        "source": Path(file_path).name,
        "file_path": str(file_path),
        "file_type": "txt",
        "page_count": None,
    }


_LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".doc": load_docx,
    ".txt": load_txt,
}


def load_document(file_path: str) -> tuple[str, dict] | None:
    ext = Path(file_path).suffix.lower()
    loader = _LOADERS.get(ext)
    if not loader:
        logger.warning("Unsupported file type: %s", ext)
        return None
    try:
        result = loader(file_path)
        logger.info("Loaded: %s", Path(file_path).name)
        return result
    except Exception as exc:
        logger.error("Error loading %s: %s", file_path, exc)
        return None


def load_directory(
    directory: str, extensions: list[str] | None = None
) -> list[tuple[str, dict]]:
    if extensions is None:
        extensions = list(SUPPORTED_EXTENSIONS)

    documents = []
    for ext in extensions:
        for file_path in Path(directory).rglob(f"*{ext}"):
            result = load_document(str(file_path))
            if result:
                documents.append(result)

    return documents
